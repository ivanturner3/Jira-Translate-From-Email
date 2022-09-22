from gettext import find
from MainConfig import *
from jira import JIRA
from deep_translator import GoogleTranslator
from docx import Document

jira = JIRA(server=jiraServer, basic_auth=(jiraUser, jiraPass))

myself = jira.myself() #returns dictionary of user information
myself = myself['displayName'] #selects user's name to use later on

foundIssue = False

while foundIssue == False:
    issueID = 'CU-136' #input("Enter Jira Issue ID: ")
    try:
        issue = jira.issue(issueID)
    except:
        errorMessage = 'Issue does not exist, please try again.'
        print(errorMessage)
    else:
        foundIssue = True

class JiraTicket:
    def __init__(self):
        self.issueID = issue
        self.summary = jira.issue(self.issueID).fields.summary
        self.description = jira.issue(self.issueID).fields.description
        self.comments = jira.comments(self.issueID)

ticket = JiraTicket()

translatedSummary = GoogleTranslator(source='auto', target='en').translate(ticket.summary)
translatedDescription = GoogleTranslator(source='auto', target='en').translate(ticket.description)

doc = Document()
#run = doc.add_paragraph()
doc.add_heading(f"Jira Ticket ID: {ticket.issueID}", 0)
doc.add_heading("Original Text:", 1)
doc.add_paragraph(f"Title: {ticket.summary}")
doc.add_paragraph(f"Description: {ticket.description}")
for comm in ticket.comments:
    if comm.jsdPublic == True:
        public = "Reply to customer"
    if comm.jsdPublic == False:
        public = "Internal comment"
    originalComment = str(comm.body).replace("\!", "!").strip()
    # findAccountID = originalComment.find("[~accountid:")
    # if findAccountID != -1:
    #     accountID = 
    para = doc.add_paragraph()
    para.add_run(f"{public} by: {comm.author.displayName}").bold = True
    para.add_run(f"\nDate Posted: {comm.created}")
    doc.add_paragraph(f"{originalComment}")
    
doc.add_heading("Translated Text:", 1)
doc.add_paragraph(f"Title: {translatedSummary}")
doc.add_paragraph(f"Description: {translatedDescription}")
for comm in ticket.comments:
    if comm.jsdPublic == True:
        public = "Reply to customer"
    if comm.jsdPublic == False:
        public = "Internal comment"
    translatedComment = str(comm.body).replace("\!", "!").strip()
    para = doc.add_paragraph()
    para.add_run(f"{public} by: {comm.author.displayName}").bold = True
    para.add_run(f"\nDate Posted: {comm.created}")
    doc.add_paragraph(f"{GoogleTranslator(source='auto', target='en').translate(translatedComment)}")
doc.save(f"translations\\{ticket.issueID}.docx")